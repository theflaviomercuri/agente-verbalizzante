#!/usr/bin/env python3
r"""
Analizza un template DOCX ed estrae tutti i placeholder {{...}} trovati.
Utile per verificare la copertura del JSON rispetto al template e per
documentare il contratto del template quando cambia.
Uso: python scripts\inspect_template.py <path_template.docx>
"""
import re
import sys
from pathlib import Path

from docx import Document

PH_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")
SECTION_TITLE_RE = re.compile(r"^(\d+)\.\s*\{\{S(\d+)T\}\}")
ACTION_RE = re.compile(r"\{\{A(\d+)\}\}")
NOTE_RE = re.compile(r"\{\{N(\d+)\}\}")


def extract_placeholders(docx_path: Path) -> dict:
    doc = Document(str(docx_path))

    all_scalars: set[str] = set()
    repeat_tables: list[list[str]] = []
    section_slots = 0
    action_slots = 0
    note_slots = 0

    # Paragrafi del body
    for p in doc.paragraphs:
        found = PH_RE.findall(p.text)
        if not found:
            continue
        for ph in found:
            all_scalars.add(ph)
        m_sec = SECTION_TITLE_RE.match(p.text.strip())
        if m_sec:
            section_slots = max(section_slots, int(m_sec.group(2)))
        m_act = ACTION_RE.search(p.text)
        if m_act:
            action_slots = max(action_slots, int(m_act.group(1)))
        m_note = NOTE_RE.search(p.text)
        if m_note:
            note_slots = max(note_slots, int(m_note.group(1)))

    # Tabelle: raccoglie i token per tabella (riga template)
    for table in doc.tables:
        table_tokens: set[str] = set()
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    found = PH_RE.findall(p.text)
                    table_tokens.update(found)
                    all_scalars.update(found)
        if table_tokens:
            repeat_tables.append(sorted(table_tokens))

    return {
        'all_placeholders': sorted(all_scalars),
        'repeat_table_token_sets': repeat_tables,
        'section_slots_in_template': section_slots,
        'action_slots_in_template': action_slots,
        'note_slots_in_template': note_slots,
    }


def check_json_coverage(result: dict, json_path: Path) -> None:
    """Se viene passato un JSON, verifica quali placeholder non sono coperti."""
    import json

    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"\n[AVVISO] Impossibile leggere il JSON per la verifica di copertura: {e}")
        return

    # Costruisce il set di chiavi JSON disponibili (solo scalari flat)
    covered: set[str] = set()
    doc = data.get('document', {})
    meeting = data.get('meeting', {})
    sections = data.get('sections', [])
    actions = data.get('actions', [])
    notes = data.get('notes', [])

    covered.update(['TITOLO', 'NOME_DOC', 'SSU', 'REF', 'AREA', 'APP', 'CTR', 'FORN',
                    'DATA', 'HIN', 'HFI', 'LUOGO', 'OGGETTO', 'INTRO',
                    'V', 'DV', 'DESCV', 'SEZV', 'AUTORE', 'ORG_AUT',
                    'DIST', 'DDIST', 'VA', 'DA', 'NA', 'OA',
                    'PN', 'PR', 'PO', 'RDOC', 'RV', 'RD', 'GT', 'GD'])
    thematic = [s for s in sections if str(s.get('number')) != '2']
    for idx in range(1, len(thematic) + 1):
        covered.add(f'S{idx}T')
        covered.add(f'S{idx}B')
    for idx in range(1, len(actions) + 1):
        covered.add(f'A{idx}')
    for idx in range(1, len(notes) + 1):
        covered.add(f'N{idx}')

    template_phs = set(result['all_placeholders'])
    uncovered = sorted(template_phs - covered)
    extra = sorted(covered - template_phs)

    print(f"\nCopertura JSON vs Template:")
    if uncovered:
        print(f"  Placeholder nel template NON coperti dal JSON ({len(uncovered)}):")
        for ph in uncovered:
            print(f"    [MANCANTE] {{{{{ph}}}}}")
    else:
        print("  Tutti i placeholder del template sono coperti dal JSON.")
    if extra:
        print(f"  Chiavi nel JSON non usate dal template ({len(extra)}):")
        for ph in extra:
            print(f"    [EXTRA]    {{{{{ph}}}}}")


def main() -> None:
    if len(sys.argv) < 2:
        print("Uso: python scripts\\inspect_template.py <path_template.docx> [path_json]")
        sys.exit(1)

    template_path = Path(sys.argv[1])
    if not template_path.exists():
        print(f"[ERRORE] File non trovato: {template_path}")
        sys.exit(1)

    result = extract_placeholders(template_path)

    print(f"\n{'=' * 55}")
    print(f"TEMPLATE: {template_path.name}")
    print(f"{'=' * 55}")

    print(f"\nPlaceholder scalari trovati ({len(result['all_placeholders'])}):")
    for ph in result['all_placeholders']:
        print(f"  {{{{{ph}}}}}")

    print(f"\nTabelle con token ripetibili ({len(result['repeat_table_token_sets'])}):")
    for i, tokens in enumerate(result['repeat_table_token_sets'], 1):
        print(f"  Tabella {i}: {tokens}")

    print(f"\nCapacita' dinamica nel template:")
    print(f"  Slot sezioni  : {result['section_slots_in_template']}")
    print(f"  Slot azioni   : {result['action_slots_in_template']}")
    print(f"  Slot note     : {result['note_slots_in_template']}")

    if len(sys.argv) >= 3:
        check_json_coverage(result, Path(sys.argv[2]))

    print(f"\n{'=' * 55}\n")


if __name__ == '__main__':
    main()
