#!/usr/bin/env python3
"""
Ricostruisce il JSON strutturato (meeting_minutes_YYYYMMDD_rev.json) a partire
dal verbale DOCX revisionato dall'utente (verbale_YYYYMMDD_v1_rev.docx).

Strategia:
  - Carica il template DOCX come riferimento strutturale per identificare
    l'indice e la struttura di ciascuna tabella (partecipanti, glossario, ecc.).
  - Legge il DOCX revisionato ed estrae il contenuto effettivo da quelle tabelle.
  - Ri-parsea le sezioni tematiche, le azioni e le note dal corpo del documento.
  - Per i metadati scalari (titolo, date, codici) riparte dal JSON originale,
    che l'utente modifica raramente in Word.

Uso:
    python scripts\\docx_reverse_map.py \\
        results\\<slug>\\verbale_YYYYMMDD_v1_rev.docx \\
        sources\\<slug>\\meeting_minutes_YYYYMMDD.json \\
        --template templates\\verbale_template_placeholders_final.docx \\
        [--output sources\\<slug>\\meeting_minutes_YYYYMMDD_rev.json]

    Se --output non è specificato, il file viene salvato automaticamente nella
    stessa directory del JSON originale con suffisso _rev.json.

Uscita:
    0  estrazione completata; file _rev.json scritto
    1  errore bloccante
"""

import argparse
import copy
import json
import re
import sys
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# Dipendenza python-docx
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.document import Document as DocType
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("ERRORE: python-docx non installata. Esegui: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# ---------------------------------------------------------------------------
# Regex
# ---------------------------------------------------------------------------
PH_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")

# Pattern per i titoli di sezione dopo il filling: "3. Titolo della sezione"
SECTION_TITLE_RE = re.compile(r"^(\d+)\.\s+(.+)$")

# Nomi sezione protetti (non tematiche)
_PROTECTED_SECTIONS = {
    "scopo del documento",
    "introduzione",
    "azioni successive",
    "note",
    "riferimenti",
    "glossario",
}

# Formato azione dopo filling: "• Owner: testo azione (Scadenza: dd/MM/yyyy)"
# oppure senza scadenza: "• Owner: testo azione"
_ACTION_LINE_RE = re.compile(
    r"^[•\-]\s*(.+?):\s*(.+?)(?:\s+\(Scadenza:\s*(\d{2}/\d{2}/\d{4})\))?$"
)

# ---------------------------------------------------------------------------
# Mapping token-set → ruolo tabella e campo-per-colonna
# ---------------------------------------------------------------------------
_TABLE_ROLES: dict[frozenset, tuple[str, tuple[str, ...]]] = {
    frozenset({"V", "DV", "DESCV", "SEZV"}):  ("history",      ("version", "date", "description", "sections")),
    frozenset({"AUTORE", "ORG_AUT"}):          ("author",       ("version", "author_name", "author_org")),
    frozenset({"DIST", "DDIST"}):              ("distribution", ("name", "organization")),
    frozenset({"VA", "DA", "NA", "OA"}):       ("approvals",    ("version", "approval_date", "name", "organization")),
    frozenset({"PN", "PR", "PO"}):             ("participants", ("name", "role", "organization")),
    frozenset({"RDOC", "RV", "RD"}):           ("references",   ("document", "version", "date")),
    frozenset({"GT", "GD"}):                   ("glossary",     ("term", "description")),
}


# ---------------------------------------------------------------------------
# Template scanning
# ---------------------------------------------------------------------------

def _scan_template_tables(template_path: Path) -> dict[str, dict]:
    """
    Scansiona il template e ritorna, per ogni ruolo-tabella trovato:
        {role: {table_idx, header_row_count, col_order}}
    dove col_order è la lista di token nell'ordine delle colonne.
    """
    doc = Document(str(template_path))
    result: dict[str, dict] = {}

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            # Raccoglie token per colonna
            col_tokens: dict[int, str] = {}
            for c_idx, cell in enumerate(row.cells):
                found = PH_RE.findall(cell.text)
                if found:
                    col_tokens[c_idx] = found[0]

            row_token_set = frozenset(col_tokens.values())

            # Cerca corrispondenza con token set noti
            for token_set, (role, _fields) in _TABLE_ROLES.items():
                if token_set.issubset(row_token_set):
                    # Ordina le colonne per indice
                    ordered_cols = sorted(col_tokens.keys())
                    col_order = [col_tokens[c] for c in ordered_cols]
                    result[role] = {
                        "table_idx": t_idx,
                        "header_row_count": r_idx,
                        "col_order": col_order,
                    }
                    break

    return result


# ---------------------------------------------------------------------------
# Table extraction from filled DOCX
# ---------------------------------------------------------------------------

def _cell_text(cell) -> str:
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _extract_table_rows(
    doc: DocType,
    table_idx: int,
    header_row_count: int,
    col_order: list[str],
) -> list[dict]:
    """
    Estrae le righe dati (dopo le righe header) dalla tabella `table_idx`
    nel documento `doc`. Ritorna lista di dict {token: valore}.
    """
    tables = doc.tables
    if table_idx >= len(tables):
        return []

    table = tables[table_idx]
    rows = table.rows
    data_rows = rows[header_row_count:]

    result = []
    for row in data_rows:
        cells = row.cells
        row_data: dict[str, str] = {}
        for col_i, token in enumerate(col_order):
            if col_i < len(cells):
                row_data[token] = _cell_text(cells[col_i]) or "-"
            else:
                row_data[token] = "-"
        # Salta righe completamente vuote o con solo trattini
        if all(v in ("", "-") for v in row_data.values()):
            continue
        result.append(row_data)

    return result


# ---------------------------------------------------------------------------
# Body paragraph parsing
# ---------------------------------------------------------------------------

def _is_section_title(para: Paragraph) -> re.Match | None:
    """Ritorna il match se il paragrafo è un titolo di sezione numerata."""
    return SECTION_TITLE_RE.match(para.text.strip())


def _is_protected(title: str) -> bool:
    return title.lower().strip() in _PROTECTED_SECTIONS


def _extract_para_text(para: Paragraph) -> str:
    """Restituisce il testo del paragrafo, splittando eventuali soft-return in \n."""
    runs_text = []
    for run in para.runs:
        runs_text.append(run.text)
    return "".join(runs_text).strip()


def _split_body_text(text: str) -> list[str]:
    """Divide il body di una sezione (che può contenere \n) in lista di paragrafi."""
    parts = [p.strip() for p in text.splitlines() if p.strip()]
    return parts if parts else [text] if text else []


def parse_document_body(doc: DocType) -> dict[str, Any]:
    """
    Parsea il corpo del documento e ritorna:
        {
          "scope_text":    str,          # sezione 1 (Scopo del documento)
          "intro_text":    str,          # sezione 2 (Introduzione)
          "thematic":      [{number, title, paragraphs}],
          "actions":       [{owner, action, due_date, status}],
          "notes":         [str],
        }
    """
    paragraphs = doc.paragraphs

    scope_text = ""
    intro_text = ""
    thematic: list[dict] = []
    actions: list[dict] = []
    notes: list[str] = []

    current_section: dict | None = None
    current_protected: str | None = None  # "azioni" | "note" | None
    i = 0

    while i < len(paragraphs):
        para = paragraphs[i]
        raw = para.text.strip()

        if not raw:
            i += 1
            continue

        m = _is_section_title(para)
        if m:
            sec_num = int(m.group(1))
            sec_title = m.group(2).strip()

            # Salva sezione corrente se era tematica
            if current_section is not None:
                thematic.append(current_section)
                current_section = None

            current_protected = None

            title_lower = sec_title.lower()
            if sec_num == 1 or "scopo" in title_lower:
                current_protected = "scope"
            elif sec_num == 2 or "introduzione" in title_lower:
                current_protected = "intro"
            elif "azioni" in title_lower:
                current_protected = "azioni"
            elif title_lower == "note":
                current_protected = "note"
            elif _is_protected(sec_title):
                current_protected = "other"
            else:
                # Sezione tematica
                current_section = {
                    "number": str(sec_num),
                    "title": sec_title,
                    "paragraphs": [],
                }
            i += 1
            continue

        # Corpo del paragrafo (non è un titolo)
        body_text = _extract_para_text(para)
        if not body_text:
            i += 1
            continue

        if current_protected == "scope":
            scope_text = (scope_text + "\n" + body_text).strip() if scope_text else body_text
        elif current_protected == "intro":
            intro_text = (intro_text + "\n" + body_text).strip() if intro_text else body_text
        elif current_protected == "azioni":
            m_act = _ACTION_LINE_RE.match(body_text)
            if m_act:
                actions.append({
                    "owner": m_act.group(1).strip(),
                    "action": m_act.group(2).strip(),
                    "due_date": m_act.group(3) or "-",
                    "status": "open",
                })
        elif current_protected == "note":
            notes.append(body_text)
        elif current_section is not None:
            # Aggiunge al body della sezione tematica corrente
            current_section["paragraphs"].extend(_split_body_text(body_text))

        i += 1

    # Salva ultima sezione tematica
    if current_section is not None:
        thematic.append(current_section)

    return {
        "scope_text": scope_text,
        "intro_text": intro_text,
        "thematic": thematic,
        "actions": actions,
        "notes": notes,
    }


# ---------------------------------------------------------------------------
# JSON reconstruction
# ---------------------------------------------------------------------------

def _t(val: Any) -> str:
    if val is None:
        return "-"
    s = str(val).strip()
    return s if s else "-"


def rebuild_json(
    original: dict,
    template_map: dict[str, dict],
    rev_doc: DocType,
) -> dict:
    """
    Combina i metadati dal JSON originale con il contenuto estratto dal DOCX revisionato.
    """
    data = copy.deepcopy(original)

    # --- Participants ---
    if "participants" in template_map:
        info = template_map["participants"]
        rows = _extract_table_rows(rev_doc, info["table_idx"], info["header_row_count"], info["col_order"])
        if rows:
            data["meeting"]["participants"] = [
                {
                    "name": _t(r.get("PN") or r.get("name")),
                    "role": _t(r.get("PR") or r.get("role")),
                    "organization": _t(r.get("PO") or r.get("organization")),
                }
                for r in rows
            ]

    # --- Distribution ---
    if "distribution" in template_map:
        info = template_map["distribution"]
        rows = _extract_table_rows(rev_doc, info["table_idx"], info["header_row_count"], info["col_order"])
        if rows:
            data["document"]["distribution"] = [
                {
                    "name": _t(r.get("DIST") or r.get("name")),
                    "role": "-",
                    "organization": _t(r.get("DDIST") or r.get("organization")),
                }
                for r in rows
            ]

    # --- Glossary ---
    if "glossary" in template_map:
        info = template_map["glossary"]
        rows = _extract_table_rows(rev_doc, info["table_idx"], info["header_row_count"], info["col_order"])
        if rows:
            data["glossary"] = [
                {
                    "term": _t(r.get("GT") or r.get("term")),
                    "description": _t(r.get("GD") or r.get("description")),
                }
                for r in rows
                if (r.get("GT") or r.get("term", "")).strip() not in ("", "-")
            ]

    # --- References ---
    if "references" in template_map:
        info = template_map["references"]
        rows = _extract_table_rows(rev_doc, info["table_idx"], info["header_row_count"], info["col_order"])
        if rows:
            non_empty = [r for r in rows if (r.get("RDOC") or r.get("document", "")).strip() not in ("", "-")]
            data["references"] = [
                {
                    "document": _t(r.get("RDOC") or r.get("document")),
                    "version": _t(r.get("RV") or r.get("version")),
                    "date": _t(r.get("RD") or r.get("date")),
                }
                for r in non_empty
            ]

    # --- History ---
    if "history" in template_map:
        info = template_map["history"]
        rows = _extract_table_rows(rev_doc, info["table_idx"], info["header_row_count"], info["col_order"])
        if rows:
            data["document"]["history"] = [
                {
                    "version": _t(r.get("V") or r.get("version")),
                    "date": _t(r.get("DV") or r.get("date")),
                    "description": _t(r.get("DESCV") or r.get("description")),
                    "sections": _t(r.get("SEZV") or r.get("sections")),
                }
                for r in rows
            ]

    # --- Body: sections, actions, notes ---
    body = parse_document_body(rev_doc)

    # Sezione 1 — Scopo del documento
    if body["scope_text"]:
        orig_sections = data.get("sections", [])
        for sec in orig_sections:
            if str(sec.get("number")) == "1":
                sec["paragraphs"] = _split_body_text(body["scope_text"])
                break
        else:
            data.setdefault("sections", []).insert(0, {
                "number": "1",
                "title": "Scopo del documento",
                "paragraphs": _split_body_text(body["scope_text"]),
            })

    # Sezione 2 — Introduzione
    if body["intro_text"]:
        orig_sections = data.get("sections", [])
        for sec in orig_sections:
            if str(sec.get("number")) == "2":
                sec["paragraphs"] = _split_body_text(body["intro_text"])
                break
        else:
            intro_entry = {
                "number": "2",
                "title": "Introduzione",
                "paragraphs": _split_body_text(body["intro_text"]),
            }
            secs = data.get("sections", [])
            data["sections"] = (
                [secs[0], intro_entry] + secs[1:] if secs else [intro_entry]
            )

    # Sezioni tematiche (3+)
    if body["thematic"]:
        # Mantieni sezioni 1 e 2 originali se non rilevate nel DOCX revisionato
        keep_12 = [s for s in data.get("sections", []) if str(s.get("number")) in ("1", "2")]
        data["sections"] = keep_12 + [
            {
                "number": sec["number"],
                "title": sec["title"],
                "paragraphs": sec["paragraphs"],
            }
            for sec in body["thematic"]
        ]

    # Actions
    if body["actions"]:
        data["actions"] = body["actions"]

    # Notes
    if body["notes"]:
        data["notes"] = body["notes"]

    return data


# ---------------------------------------------------------------------------
# Entrypoint
# ---------------------------------------------------------------------------

def main() -> None:
    ap = argparse.ArgumentParser(
        description=(
            "Ricostruisce il JSON _rev dal verbale DOCX revisionato dall'utente."
        )
    )
    ap.add_argument(
        "rev_docx",
        help="Verbale DOCX revisionato (es. results/<slug>/verbale_YYYYMMDD_v1_rev.docx)",
    )
    ap.add_argument(
        "original_json",
        help="JSON originale generato dall'agente (es. sources/<slug>/meeting_minutes_YYYYMMDD.json)",
    )
    ap.add_argument(
        "--template",
        required=True,
        help="Template DOCX con placeholder (es. templates/verbale_template_placeholders_final.docx)",
    )
    ap.add_argument(
        "--output",
        default=None,
        help=(
            "Percorso di output per il JSON revisionato. "
            "Default: accanto all'original_json con suffisso _rev.json"
        ),
    )
    args = ap.parse_args()

    rev_docx_path = Path(args.rev_docx)
    original_json_path = Path(args.original_json)
    template_path = Path(args.template)

    # Validazioni
    for p, label in (
        (rev_docx_path, "rev_docx"),
        (original_json_path, "original_json"),
        (template_path, "template"),
    ):
        if not p.exists():
            print(f"ERRORE: file non trovato ({label}): {p}", file=sys.stderr)
            sys.exit(1)

    # Output path
    if args.output:
        output_path = Path(args.output)
    else:
        stem = original_json_path.stem
        # Rimuovi eventuali suffissi _rev esistenti prima di aggiungerne uno nuovo
        base_stem = stem.removesuffix("_rev")
        output_path = original_json_path.parent / f"{base_stem}_rev.json"

    # Caricamento
    print(f"Template: {template_path}", file=sys.stderr)
    template_map = _scan_template_tables(template_path)
    print(
        f"Tabelle identificate: {', '.join(template_map.keys()) or 'nessuna'}",
        file=sys.stderr,
    )

    with open(original_json_path, encoding="utf-8") as f:
        original = json.load(f)

    rev_doc = Document(str(rev_docx_path))

    # Ricostruzione
    rev_json = rebuild_json(original, template_map, rev_doc)

    # Scrittura
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(rev_json, f, ensure_ascii=False, indent=2)

    print(f"Scritto: {output_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
