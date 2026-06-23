import argparse
import copy
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence, Tuple

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph

PH_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")
SECTION_TITLE_RE = re.compile(r"^(\d+)\.\s*\{\{S(\d+)T\}\}\s*$")
ACTION_RE = re.compile(r"\{\{A(\d+)\}\}")
NOTE_RE = re.compile(r"\{\{N(\d+)\}\}")
_PROTECTED_TITLE_RE = re.compile(
    r'^(?:\d+\.\s*)?(Azioni successive|Note|Scopo del documento|Introduzione|'
    r'Dati di sintesi|Riferimenti|Glossario)$',
    re.IGNORECASE,
)
_RENUMBER_RE = re.compile(r'^(\d+)\.\s*(Azioni successive|Note)\s*$', re.IGNORECASE)


def load_json(path: str | Path) -> Dict[str, Any]:
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def fmt_date(val: Any) -> str:
    if not val or val == '-':
        return '-'
    if isinstance(val, str):
        v = val.strip()
        for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y'):
            try:
                return datetime.strptime(v, fmt).strftime('%d/%m/%Y')
            except ValueError:
                pass
        return v
    return str(val)


def text(val: Any) -> str:
    if val is None:
        return '-'
    if isinstance(val, list):
        vals = [text(v) for v in val if str(v).strip()]
        return ', '.join(vals) if vals else '-'
    s = str(val).strip()
    return s if s else '-'


def join_paragraphs(paragraphs: Sequence[str]) -> str:
    cleaned = [p.strip() for p in paragraphs if p and p.strip()]
    return '\n'.join(cleaned) if cleaned else '-'


def iter_all_paragraphs(doc: DocumentObject) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def normalize_runs(paragraph: Paragraph) -> None:
    """Merge all runs into the first one when the paragraph contains a placeholder.
    Prevents placeholders fragmented across runs (by Word spell-check or style
    tracking) from being missed or incorrectly replaced."""
    if not PH_RE.search(paragraph.text):
        return
    if len(paragraph.runs) <= 1:
        return
    full_text = ''.join(run.text for run in paragraph.runs)
    paragraph.runs[0].text = full_text
    for run in paragraph.runs[1:]:
        run.text = ''


def replace_in_paragraph(paragraph: Paragraph, mapping: Dict[str, str]) -> None:
    normalize_runs(paragraph)
    original = paragraph.text
    updated = original
    for key, value in mapping.items():
        updated = updated.replace('{{' + key + '}}', value)
    if updated == original:
        return
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(updated)


def replace_everywhere(doc: DocumentObject, mapping: Dict[str, str]) -> None:
    for p in iter_all_paragraphs(doc):
        replace_in_paragraph(p, mapping)


def find_table_row_by_tokens(doc: DocumentObject, tokens: Sequence[str]) -> Tuple[Table, int]:
    token_set = set(tokens)
    for table in doc.tables:
        for idx, row in enumerate(table.rows):
            row_tokens = set()
            for cell in row.cells:
                row_tokens.update(PH_RE.findall(cell.text))
            if token_set.issubset(row_tokens):
                return table, idx
    raise ValueError(f'Riga template non trovata per token {tokens}')


def clone_row_after_xml(table: Table, anchor_row_idx: int, template_tr) -> _Row:
    new_tr = copy.deepcopy(template_tr)
    table.rows[anchor_row_idx]._tr.addnext(new_tr)
    return _Row(new_tr, table)


def remove_row(table: Table, row_idx: int) -> None:
    tr = table.rows[row_idx]._tr
    table._tbl.remove(tr)


def replace_in_row(row: _Row, mapping: Dict[str, str]) -> None:
    for cell in row.cells:
        for p in cell.paragraphs:
            replace_in_paragraph(p, mapping)


def populate_repeat_table(doc: DocumentObject, tokens: Sequence[str], rows_data: Sequence[Sequence[str]]) -> None:
    table, row_idx = find_table_row_by_tokens(doc, tokens)
    if not rows_data:
        rows_data = [['-'] * len(tokens)]
    template_tr = copy.deepcopy(table.rows[row_idx]._tr)
    insert_at = row_idx
    for values in rows_data:
        new_row = clone_row_after_xml(table, insert_at, template_tr)
        replace_in_row(new_row, dict(zip(tokens, [text(v) for v in values])))
        insert_at += 1
    remove_row(table, row_idx)


def build_scalar_mapping(data: Dict[str, Any]) -> Dict[str, str]:
    d = data.get('document', {})
    m = data.get('meeting', {})
    history = (d.get('history') or [{
        'version': d.get('version', '1.0'),
        'date': m.get('date') or '-',
        'description': 'Versione iniziale',
        'sections': ''
    }])[0]
    author = d.get('author', {}) or {}
    distribution = (d.get('distribution') or [{'list': '-', 'date': '-'}])[0]
    approvals = (d.get('approvals') or [{'version': d.get('version', '-'), 'approval_date': '-', 'name': '-', 'organization': '-'}])[0]

    sections = data.get('sections', [])
    intro_section = next((s for s in sections if str(s.get('number')) == '2'), {})
    thematic = [s for s in sections if str(s.get('number')) not in ('1', '2')]

    actions = data.get('actions', [])
    notes = data.get('notes', [])

    mapping = {
        'TITOLO': text(d.get('title', m.get('subject', '-'))),
        'NOME_DOC': text(d.get('document_name')),
        'SSU': text(d.get('ssu_code')),
        'AREA': text(d.get('management_area')),
        'APP': text(d.get('application_code')),
        'CTR': text(d.get('contract_reference')),
        'FORN': text(d.get('supplier')),
        'DATA': fmt_date(m.get('date')),
        'HIN': text(m.get('start_time')),
        'HFI': text(m.get('end_time')),
        'LUOGO': text(m.get('location')),
        'OGGETTO': text(m.get('subject')),
        'INTRO': join_paragraphs(intro_section.get('paragraphs', [])),
    }

    for idx, sec in enumerate(thematic, start=1):
        mapping[f'S{idx}T'] = text(sec.get('title'))
        mapping[f'S{idx}B'] = join_paragraphs(sec.get('paragraphs', []))

    # REF auto-derivato dagli INPS participants se client_references è vuoto
    client_refs = d.get('client_references') or []
    if client_refs:
        ref_str = text(client_refs)
    else:
        inps_parts = [p.get('name', '') for p in (m.get('participants') or []) if 'INPS' in str(p.get('organization', ''))]
        ref_str = ', '.join(inps_parts) if inps_parts else '-'
    mapping['REF'] = ref_str

    for idx, act in enumerate(actions, start=1):
        owner = text(act.get('owner'))
        action_text_val = text(act.get('action'))
        label = f'{owner}: {action_text_val}'
        mapping[f'A{idx}'] = label

    for idx, note in enumerate(notes, start=1):
        mapping[f'N{idx}'] = text(note)

    return mapping


def build_repeat_data(data: Dict[str, Any]) -> Dict[Tuple[str, ...], List[List[str]]]:
    doc = data.get('document', {})
    meeting = data.get('meeting', {})
    history_rows = doc.get('history') or [{
        'version': doc.get('version', '1.0'),
        'date': meeting.get('date') or '-',
        'description': 'Versione iniziale',
        'sections': ''
    }]
    distribution_rows = doc.get('distribution') or [{'list': '-', 'date': '-'}]
    approval_rows = doc.get('approvals') or [{'version': doc.get('version', '-'), 'approval_date': '-', 'name': '-', 'organization': '-'}]
    participants = meeting.get('participants', []) or []
    references = data.get('references', []) or []
    glossary = data.get('glossary', []) or []

    return {
        ('V', 'DV', 'DESCV', 'SEZV'): [[text(r.get('version')), fmt_date(r.get('date')), text(r.get('description')), text(r.get('sections'))] for r in history_rows],
        ('V', 'AUTORE', 'ORG_AUT'): [[text(doc.get('version', '1.0')), text(doc.get('author', {}).get('name')), text(doc.get('author', {}).get('organization'))]],
        ('DIST', 'DDIST'): [[text(r.get('name')), text(r.get('organization'))] for r in distribution_rows],
        ('VA', 'DA', 'NA', 'OA'): [[text(r.get('version')), fmt_date(r.get('approval_date')), text(r.get('name')), text(r.get('organization'))] for r in approval_rows],
        ('PN', 'PR', 'PO'): [[text(p.get('name')), text(p.get('role')), text(p.get('organization'))] for p in participants],
        ('RDOC', 'RV', 'RD'): [[text(r.get('document')), text(r.get('version')), fmt_date(r.get('date'))] for r in references],
        ('GT', 'GD'): [[text(g.get('term') or g.get('acronym')), text(g.get('description'))] for g in glossary],
    }


def insert_paragraph_after(paragraph: Paragraph, text_value: str = '') -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    try:
        new_para.style = paragraph.style
    except Exception:
        pass
    if text_value:
        new_para.add_run(text_value)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def ensure_dynamic_sections(doc: DocumentObject, data: Dict[str, Any]) -> None:
    sections = [s for s in data.get('sections', []) if str(s.get('number')) not in ('1', '2')]
    title_placeholders = []
    paras = doc.paragraphs
    for p in paras:
        m = SECTION_TITLE_RE.match(p.text.strip())
        if m:
            title_placeholders.append((int(m.group(1)), int(m.group(2)), p))
    title_placeholders.sort(key=lambda x: x[1])
    capacity = len(title_placeholders)
    if not title_placeholders or len(sections) <= capacity:
        return
    last_num, _, last_title_para = title_placeholders[-1]
    title_idx = paras.index(last_title_para)
    body_para = paras[title_idx + 1]
    cursor = body_para
    current_num = last_num
    for sec in sections[capacity:]:
        current_num += 1
        new_title = insert_paragraph_after(cursor, f'{current_num}. {text(sec.get("title"))}')
        try:
            new_title.style = last_title_para.style
        except Exception:
            pass
        new_body = insert_paragraph_after(new_title, join_paragraphs(sec.get('paragraphs', [])))
        try:
            new_body.style = body_para.style
        except Exception:
            pass
        cursor = new_body


def ensure_dynamic_actions(doc: DocumentObject, data: Dict[str, Any]) -> None:
    actions = data.get('actions', [])
    action_paras = []
    for p in doc.paragraphs:
        m = ACTION_RE.search(p.text)
        if m:
            action_paras.append((int(m.group(1)), p))
    action_paras.sort(key=lambda x: x[0])
    capacity = len(action_paras)
    if not action_paras or len(actions) <= capacity:
        return
    cursor = action_paras[-1][1]
    for action in actions[capacity:]:
        owner = text(action.get('owner'))
        action_text = text(action.get('action'))
        line = f'• {owner}: {action_text}'
        new_p = insert_paragraph_after(cursor, line)
        try:
            new_p.style = cursor.style
        except Exception:
            pass
        cursor = new_p


def ensure_dynamic_notes(doc: DocumentObject, data: Dict[str, Any]) -> None:
    notes = data.get('notes', [])
    note_paras = []
    for p in doc.paragraphs:
        m = NOTE_RE.search(p.text)
        if m:
            note_paras.append((int(m.group(1)), p))
    note_paras.sort(key=lambda x: x[0])
    capacity = len(note_paras)
    if not note_paras or len(notes) <= capacity:
        return
    cursor = note_paras[-1][1]
    for note in notes[capacity:]:
        new_p = insert_paragraph_after(cursor, text(note))
        try:
            new_p.style = cursor.style
        except Exception:
            pass
        cursor = new_p


def insert_logo_first_page(doc: DocumentObject, logo_path: Path) -> None:
    """Inserisce il logo centrato come primo paragrafo del documento."""
    if not logo_path.exists():
        print(f'WARN: logo non trovato: {logo_path}', file=sys.stderr)
        return
    first_para = doc.paragraphs[0]
    new_p = OxmlElement('w:p')
    first_para._p.addprevious(new_p)
    logo_para = Paragraph(new_p, first_para._parent)
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_para.add_run()
    run.add_picture(str(logo_path), width=Inches(2.0))


def renumber_protected_sections(doc: DocumentObject, data: Dict[str, Any]) -> None:
    """Rinumera progressivamente le sezioni 'Azioni successive' e 'Note'
    in modo che la loro numerazione segua le sezioni tematiche."""
    thematic = [s for s in data.get('sections', []) if str(s.get('number')) not in ('1', '2')]
    # 1 = Scopo, 2 = Introduzione, poi N sezioni tematiche
    next_num = 2 + len(thematic) + 1
    for p in doc.paragraphs:
        m = _RENUMBER_RE.match(p.text.strip())
        if m:
            label = m.group(2)
            if p.runs:
                p.runs[0].text = f'{next_num}. {label}'
                for r in p.runs[1:]:
                    r.text = ''
            else:
                p.add_run(f'{next_num}. {label}')
            next_num += 1


def prune_empty_paragraphs(doc: DocumentObject) -> None:
    for p in list(doc.paragraphs):
        raw = p.text.strip()
        if _PROTECTED_TITLE_RE.match(raw):
            continue
        cleaned = raw.replace('•', '').replace('-', '').strip()
        if cleaned:
            continue
        remove_paragraph(p)


def fill_template(template_path: str | Path, json_path: str | Path, output_path: str | Path, logo_path: str | Path | None = None) -> None:
    data = load_json(json_path)
    doc = Document(str(template_path))
    # Logo INPS centrato in cima alla pagina 1
    if logo_path is None:
        logo_path = Path(template_path).parent / 'logo inps.png'
    insert_logo_first_page(doc, Path(logo_path))
    for tokens, rows in build_repeat_data(data).items():
        populate_repeat_table(doc, list(tokens), rows)
    ensure_dynamic_sections(doc, data)
    ensure_dynamic_actions(doc, data)
    ensure_dynamic_notes(doc, data)
    replace_everywhere(doc, build_scalar_mapping(data))
    # Rinumera "Azioni successive" e "Note" dopo le sezioni tematiche
    renumber_protected_sections(doc, data)
    remaining = sorted(set(PH_RE.findall(' '.join(p.text for p in iter_all_paragraphs(doc)))))
    if remaining:
        for k in remaining:
            print(f'WARN: placeholder non riempito rimosso: {{{{{k}}}}}', file=sys.stderr)
    replace_everywhere(doc, {k: '' for k in remaining})
    prune_empty_paragraphs(doc)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    ap = argparse.ArgumentParser(description='Compila un template DOCX con placeholder reali a partire da JSON.')
    ap.add_argument('input_json')
    ap.add_argument('output_docx')
    ap.add_argument('--template', required=True)
    ap.add_argument('--logo', default=None, help='Percorso al logo (default: logo inps.png nella cartella del template)')
    args = ap.parse_args()
    fill_template(args.template, args.input_json, args.output_docx, logo_path=args.logo)


if __name__ == '__main__':
    main()
